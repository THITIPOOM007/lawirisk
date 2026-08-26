from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = Path(r"D:\LAWiRISK project lew\output\documents\ร่างโครงการ_LawiRisk-SSK_Evidence_Intelligence_ขีดเส้นใต้ส่วนเพิ่มเติม.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)

FONT = "Angsana New"
BODY_SIZE = 16
ACCENT = RGBColor(31, 77, 120)
MUTED = RGBColor(92, 92, 92)
UNDERLINE_COLOR = RGBColor(0, 0, 0)


def set_run_font(run, size=BODY_SIZE, bold=False, italic=False, underline=False, color=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    if color:
        run.font.color.rgb = color
    return run


def set_para_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=6, line=1.25,
                    first_line=0.5, keep_with_next=False):
    pf = p.paragraph_format
    pf.alignment = align
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if first_line is not None:
        pf.first_line_indent = Inches(first_line)
    pf.keep_with_next = keep_with_next


def add_segments(doc, segments, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0,
                 after=6, line=1.25, first_line=0.5, keep_with_next=False,
                 style=None):
    p = doc.add_paragraph(style=style)
    set_para_format(p, align, before, after, line, first_line, keep_with_next)
    for seg in segments:
        if isinstance(seg, str):
            seg = {"text": seg}
        run = p.add_run(seg["text"])
        set_run_font(
            run,
            size=seg.get("size", BODY_SIZE),
            bold=seg.get("bold", False),
            italic=seg.get("italic", False),
            underline=seg.get("underline", False),
            color=seg.get("color"),
        )
    return p


def add_plain(doc, text, **kwargs):
    return add_segments(doc, [{"text": text}], **kwargs)


def add_under(doc, text, **kwargs):
    return add_segments(doc, [{"text": text, "underline": True}], **kwargs)


def add_section_heading(doc, number, title):
    p = doc.add_paragraph(style="Heading 1")
    set_para_format(p, WD_ALIGN_PARAGRAPH.LEFT, before=10, after=5, line=1.0,
                    first_line=None, keep_with_next=True)
    set_run_font(p.add_run(f"{number}. {title}"), size=18, bold=True, color=ACCENT)
    return p


def add_subheading(doc, label, title, under=False):
    p = doc.add_paragraph(style="Heading 2")
    set_para_format(p, WD_ALIGN_PARAGRAPH.LEFT, before=7, after=3, line=1.0,
                    first_line=None, keep_with_next=True)
    set_run_font(p.add_run(f"{label} {title}"), size=16, bold=True,
                 underline=under, color=RGBColor(0, 0, 0))
    return p


def new_numbering_group(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = (max(abstract_ids) + 1) if abstract_ids else 1
    num_id = (max(num_ids) + 1) if num_ids else 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    ppr.append(tabs)
    ppr.append(ind)
    lvl.extend([start, num_fmt, lvl_text, suff, ppr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    aid = OxmlElement("w:abstractNumId")
    aid.set(qn("w:val"), str(abstract_id))
    num.append(aid)
    numbering.append(num)
    return num_id


def add_number_item(doc, text, *, num_id, under=False, bold_prefix=None):
    p = doc.add_paragraph(style="Normal")
    ppr = p._p.get_or_add_pPr()
    numpr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(num_id))
    numpr.extend([ilvl, numid])
    ppr.append(numpr)
    pf = p.paragraph_format
    pf.left_indent = Inches(0.5)
    pf.first_line_indent = Inches(-0.25)
    pf.space_before = Pt(0)
    pf.space_after = Pt(4)
    pf.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        set_run_font(p.add_run(bold_prefix), bold=True, underline=under)
        set_run_font(p.add_run(text[len(bold_prefix):]), underline=under)
    else:
        set_run_font(p.add_run(text), underline=under)
    return p


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def set_row_cant_split(row):
    trPr = row._tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    trPr.append(cant)


def set_cell_width(cell, dxa):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(dxa))
    tcW.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(widths)))
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "120")
    tblInd.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[i])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_text(cell, text, *, bold=False, underline=False, align=WD_ALIGN_PARAGRAPH.LEFT,
                  size=14, color=None):
    p = cell.paragraphs[0]
    p.clear()
    set_para_format(p, align, before=0, after=0, line=1.15, first_line=None)
    set_run_font(p.add_run(text), size=size, bold=bold, underline=underline, color=color)


def add_budget_table(doc):
    rows = [
        ("1", "กิจกรรมประชุมคณะทำงานกำหนดความต้องการและทบทวนระบบ (1 ครั้ง 20 คน)", "2,800"),
        ("1.1", "ค่าอาหารกลางวัน (20 คน x 80 บาท x 1 มื้อ)", "1,600"),
        ("1.2", "ค่าอาหารว่างและเครื่องดื่ม (20 คน x 30 บาท x 2 มื้อ)", "1,200"),
        ("2", "กิจกรรมอบรมเชิงปฏิบัติการใช้งานระบบ (1 ครั้ง 120 คน)", "19,800"),
        ("2.1", "ค่าอาหารกลางวัน (120 คน x 80 บาท x 1 มื้อ)", "9,600"),
        ("2.2", "ค่าอาหารว่างและเครื่องดื่ม (120 คน x 30 บาท x 1 มื้อ)", "3,600"),
        ("2.3", "ค่าสมนาคุณวิทยากรบรรยายหลัก (2 คน x 600 บาท x 3 ชั่วโมง)", "3,600"),
        ("2.4", "ค่าพาหนะจ่ายจริงสำหรับ อสม./เครือข่ายภาคประชาชน (30 คน x 100 บาท)", "3,000"),
        ("", "รวมเป็นเงินงบประมาณทั้งสิ้น", "22,600"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(("รายการ", "รายละเอียด", "จำนวนเงิน (บาท)")):
        shade_cell(hdr[i], "E8EEF5")
        set_cell_text(hdr[i], text, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=ACCENT)
    set_repeat_table_header(table.rows[0])
    for idx, desc, amount in rows:
        cells = table.add_row().cells
        set_row_cant_split(table.rows[-1])
        is_total = idx == ""
        set_cell_text(cells[0], idx, align=WD_ALIGN_PARAGRAPH.CENTER, bold=is_total)
        set_cell_text(cells[1], desc, bold=is_total, underline=True)
        set_cell_text(cells[2], amount, align=WD_ALIGN_PARAGRAPH.RIGHT, bold=is_total)
        if idx in ("1", "2"):
            for c in cells:
                shade_cell(c, "F4F6F9")
        if is_total:
            for c in cells:
                shade_cell(c, "DDEBF7")
    set_table_geometry(table, [1200, 6560, 1600])
    doc.add_paragraph()
    return table


def add_schedule_table(doc):
    data = [
        ("สิงหาคม 2569", "แต่งตั้งคณะทำงาน วิเคราะห์กระบวนงาน จัดทำข้อกำหนดข้อมูลและสิทธิผู้ใช้"),
        ("สิงหาคม-กันยายน 2569", "ปรับปรุงระบบ จัดทำคู่มือและมาตรฐานการจัดการพยานหลักฐาน ทดสอบระบบและความปลอดภัย"),
        ("กันยายน 2569", "ทดสอบการยอมรับโดยผู้ใช้ (UAT) อบรม และเปิดใช้งานนำร่อง"),
        ("ตุลาคม-ธันวาคม 2569", "ติดตามผล ประเมินตัวชี้วัด แก้ไขข้อบกพร่อง และจัดทำรายงาน โดยใช้ทรัพยากรประจำของหน่วยงาน"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, text in enumerate(("ช่วงเวลา", "กิจกรรมสำคัญ")):
        shade_cell(table.rows[0].cells[i], "E8EEF5")
        set_cell_text(table.rows[0].cells[i], text, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=ACCENT)
    set_repeat_table_header(table.rows[0])
    for period, activity in data:
        cells = table.add_row().cells
        set_cell_text(cells[0], period, underline=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(cells[1], activity, underline=True)
    set_table_geometry(table, [2200, 7160])
    doc.add_paragraph()


doc = Document()
doc.settings.odd_and_even_pages_header_footer = False
sec = doc.sections[0]
sec.page_width = Cm(21.0)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.2)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.0)
sec.header_distance = Cm(1.2)
sec.footer_distance = Cm(1.2)

# Base styles: narrative_proposal with named Thai government A4/type override.
normal = doc.styles["Normal"]
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
normal.font.size = Pt(BODY_SIZE)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25
for style_name, size, color in (("Heading 1", 18, ACCENT), ("Heading 2", 16, RGBColor(0,0,0))):
    st = doc.styles[style_name]
    st.font.name = FONT
    st._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    st._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = color

# Quiet footer only; government project draft uses no running header.
fp = sec.footer.paragraphs[0]
set_para_format(fp, WD_ALIGN_PARAGRAPH.CENTER, after=0, line=1.0, first_line=None)
set_run_font(fp.add_run("หน้า "), size=10, color=MUTED)
fld = OxmlElement("w:fldSimple")
fld.set(qn("w:instr"), "PAGE")
fp._p.append(fld)

# Proposal centerpiece.
add_plain(doc, "สำนักงานสาธารณสุขจังหวัดศรีสะเกษ", align=WD_ALIGN_PARAGRAPH.CENTER,
          first_line=None, after=2, line=1.0)
p = doc.add_paragraph()
set_para_format(p, WD_ALIGN_PARAGRAPH.CENTER, before=14, after=4, line=1.0, first_line=None)
set_run_font(p.add_run("ร่างโครงการ"), size=24, bold=True, color=ACCENT)
p = doc.add_paragraph()
set_para_format(p, WD_ALIGN_PARAGRAPH.CENTER, after=3, line=1.05, first_line=None)
set_run_font(p.add_run("พัฒนาและนำระบบ LawiRisk-SSK Evidence Intelligence มาใช้ในการสืบสวน วิเคราะห์ความเชื่อมโยง และบริหารจัดการพยานหลักฐานดิจิทัล"), size=20, bold=True, underline=True)
add_plain(doc, "ประจำปีงบประมาณ พ.ศ. 2569", align=WD_ALIGN_PARAGRAPH.CENTER,
          first_line=None, after=10, line=1.0)
add_segments(doc, [
    {"text": "หมายเหตุ: ", "bold": True, "size": 14},
    {"text": "ข้อความที่ขีดเส้นใต้เป็นข้อความที่เสนอให้เพิ่มหรือแก้ไขจากร่างเดิม", "underline": True, "size": 14},
], align=WD_ALIGN_PARAGRAPH.CENTER, first_line=None, after=16, line=1.0)

add_section_heading(doc, "1", "ชื่อโครงการ")
add_under(doc, "โครงการพัฒนาและนำระบบ LawiRisk-SSK Evidence Intelligence มาใช้ในการสืบสวน วิเคราะห์ความเชื่อมโยง และบริหารจัดการพยานหลักฐานดิจิทัล สำนักงานสาธารณสุขจังหวัดศรีสะเกษ ประจำปีงบประมาณ พ.ศ. 2569")

add_section_heading(doc, "2", "หลักการและเหตุผล")
add_plain(doc, "จากสถานการณ์เรื่องร้องเรียนด้านผลิตภัณฑ์และบริการสุขภาพในปัจจุบัน พบว่ามีแนวโน้มเพิ่มสูงขึ้น โดยเฉพาะเรื่องร้องเรียนเกี่ยวกับอาหาร อาหารเสริม ยา เครื่องสำอาง สถานพยาบาล คลินิก และสถานประกอบการเพื่อสุขภาพ ประกอบกับการจำหน่ายผลิตภัณฑ์และบริการผ่านช่องทางออนไลน์ทำให้รูปแบบการกระทำผิดมีความซับซ้อนและเปลี่ยนแปลงอย่างรวดเร็ว")
add_plain(doc, "จังหวัดศรีสะเกษมีเรื่องร้องเรียนเกี่ยวกับผลิตภัณฑ์และบริการสุขภาพไม่น้อยกว่า 50 เรื่องต่อปี ขณะที่พนักงานเจ้าหน้าที่ผู้รับผิดชอบงานกฎหมายและการดำเนินคดีมีอัตรากำลังไม่สอดคล้องกับภาระงาน การตรวจสอบพยานหลักฐานดิจิทัลในแต่ละเรื่องต้องใช้เวลาเฉลี่ยประมาณ 5-8 ชั่วโมง และข้อมูลสำคัญกระจายอยู่ในหลายรูปแบบ เช่น ภาพหน้าจอ URL หมายเลขโทรศัพท์ บัญชีธนาคาร ที่อยู่จัดส่ง เพจหรือบัญชีสื่อสังคมออนไลน์ และเอกสารประกอบคดี")
add_plain(doc, "ผู้กระทำผิดอาจสร้างเครือข่ายเพจหรือบัญชีอวตาร ใช้ชื่อหน้าร้านหลายชื่อ และเปลี่ยนช่องทางติดต่อเพื่อหลีกเลี่ยงการตรวจสอบ แม้หน้าร้านจะแตกต่างกัน แต่เบื้องหลังอาจมีจุดเชื่อมโยงร่วมกัน การตรวจสอบด้วยวิธีเดิมจึงใช้เวลามากและยากต่อการเชื่อมโยงข้อมูลข้ามเรื่องร้องเรียนหรือคดีในอดีต")
add_under(doc, "สำนักงานสาธารณสุขจังหวัดศรีสะเกษจึงได้พัฒนา LawiRisk-SSK Evidence Intelligence เป็นระบบสนับสนุนการปฏิบัติงานสืบสวนและบริหารจัดการพยานหลักฐานดิจิทัล โดยมุ่งเน้นการรักษาความสมบูรณ์ของหลักฐาน การบันทึกแหล่งที่มาและประวัติการครอบครองหลักฐาน การจำกัดสิทธิการเข้าถึงรายคดี การวิเคราะห์ความเชื่อมโยงที่ตรวจสอบย้อนกลับได้ และการช่วยจัดทำร่างเอกสารทางกฎหมาย")
add_under(doc, "ระบบดังกล่าวเป็นเพียงเครื่องมือสนับสนุนการวิเคราะห์ มิได้ใช้ตัดสินข้อเท็จจริง ข้อกฎหมาย ความผิด หรือผลทางคดีโดยอัตโนมัติ ผลวิเคราะห์และร่างเอกสารทุกฉบับต้องผ่านการตรวจสอบและรับรองโดยพนักงานเจ้าหน้าที่ผู้มีอำนาจก่อนนำไปใช้ในการดำเนินงาน")
add_under(doc, "เพื่อให้การนำระบบไปใช้จริงมีมาตรฐานเดียวกัน จึงจำเป็นต้องกำหนดกระบวนงาน บทบาทผู้ใช้งาน มาตรการคุ้มครองข้อมูลส่วนบุคคล มาตรฐานการจัดการพยานหลักฐานดิจิทัล เกณฑ์ทดสอบและรับมอบระบบ ตลอดจนแผนบำรุงรักษาและติดตามผลอย่างชัดเจน")

add_section_heading(doc, "3", "วัตถุประสงค์")
num_obj = new_numbering_group(doc)
add_number_item(doc, "เพื่อพัฒนาระบบสืบค้นและวิเคราะห์ความเชื่อมโยงพยานหลักฐานดิจิทัล รวมถึงช่วยจัดทำร่างแบบฟอร์มและเอกสารทางกฎหมาย", num_id=num_obj)
add_number_item(doc, "เพื่อพัฒนาศักยภาพพนักงานเจ้าหน้าที่ในการประยุกต์ใช้เทคโนโลยีดิจิทัลเพื่อการสืบสวน รวบรวมพยานหลักฐาน และบังคับใช้กฎหมายคุ้มครองผู้บริโภคได้อย่างรวดเร็วและมีประสิทธิภาพ", num_id=num_obj)
add_number_item(doc, "เพื่อกำหนดมาตรฐานการรวบรวม เก็บรักษา ส่งมอบ ตรวจสอบ และส่งออกพยานหลักฐานดิจิทัลให้สามารถตรวจสอบย้อนกลับได้ตลอดวงจรชีวิตของหลักฐาน", num_id=num_obj, under=True)
add_number_item(doc, "เพื่อจัดให้มีการควบคุมสิทธิการเข้าถึง การบันทึกประวัติการใช้งาน การสำรองและกู้คืนข้อมูล และมาตรการคุ้มครองข้อมูลส่วนบุคคลที่เหมาะสมกับระดับความเสี่ยง", num_id=num_obj, under=True)
add_number_item(doc, "เพื่อทดลองใช้และประเมินประสิทธิผลของระบบก่อนขยายผล โดยให้คณะทำงานด้านกฎหมายตรวจสอบผลวิเคราะห์ทุกกรณีในระยะนำร่อง", num_id=num_obj, under=True)

add_section_heading(doc, "4", "กลุ่มเป้าหมาย")
num_target = new_numbering_group(doc)
add_number_item(doc, "คณะทำงานออกแบบ พัฒนา และกำกับดูแลระบบ LawiRisk-SSK Evidence Intelligence จำนวน 30 คน", num_id=num_target, under=True)
add_number_item(doc, "เจ้าหน้าที่ผู้รับผิดชอบงานคุ้มครองผู้บริโภคระดับโรงพยาบาลและโรงพยาบาลส่งเสริมสุขภาพตำบล จำนวน 50 คน", num_id=num_target)
add_number_item(doc, "คณะทำงานด้านกฎหมายและการดำเนินคดีผลิตภัณฑ์สุขภาพ สำนักงานสาธารณสุขจังหวัดศรีสะเกษ จำนวน 10 คน", num_id=num_target)
add_number_item(doc, "อสม. เชี่ยวชาญงานคุ้มครองผู้บริโภคและเครือข่ายภาคประชาชน จำนวน 30 คน โดยเข้ารับการอบรมเฉพาะการแจ้งเบาะแส การเก็บข้อมูลเบื้องต้น และการใช้บริการสำหรับประชาชน ไม่มีสิทธิเข้าถึงข้อมูลสำนวนคดี", num_id=num_target, under=True)

add_section_heading(doc, "5", "ขอบเขตและคุณลักษณะขั้นต่ำของระบบ")
num_scope = new_numbering_group(doc)
for text in [
    "ระบบรับและบริหารเรื่องร้องเรียนหรือสำนวนคดี โดยกำหนดเลขอ้างอิง สถานะ ผู้รับผิดชอบ และลำดับเหตุการณ์ของแต่ละเรื่อง",
    "ระบบจัดเก็บพยานหลักฐานต้นฉบับโดยไม่แก้ไข พร้อมบันทึกผู้เก็บ วันเวลา แหล่งที่มา URL ข้อมูลกำกับ และค่าแฮช SHA-256 หรือวิธีการอื่นที่มีความน่าเชื่อถือเทียบเท่า",
    "ระบบบันทึกประวัติการครอบครองและการดำเนินการต่อหลักฐาน (Chain of Custody) รวมถึงการรับ ส่ง เปิดดู ดาวน์โหลด ส่งออก และการเปลี่ยนแปลงข้อมูลประกอบ",
    "ระบบวิเคราะห์และแสดงความเชื่อมโยงระหว่างบุคคล เพจ หมายเลขโทรศัพท์ บัญชีธนาคาร ที่อยู่ ผู้ให้บริการขนส่ง หรือข้อมูลอื่นที่เกี่ยวข้อง โดยทุกจุดเชื่อมโยงต้องย้อนกลับไปตรวจสอบหลักฐานต้นทางได้",
    "ระบบช่วยจัดทำร่างแบบฟอร์มหรือเอกสารทางกฎหมายจากข้อมูลที่เจ้าหน้าที่เลือก โดยแสดงแหล่งข้อมูลและสถานะว่าเป็นร่างรอตรวจสอบ",
    "ระบบกำหนดสิทธิตามบทบาทและรายคดี แยกอย่างน้อยเป็นผู้รับเรื่อง ผู้สืบสวน ผู้ตรวจทานด้านกฎหมาย ผู้อนุมัติ และผู้ดูแลระบบ",
    "ระบบบันทึกเหตุการณ์การใช้งาน (Audit Log) สำรองข้อมูล ทดสอบการกู้คืน และแจ้งเตือนเหตุผิดปกติที่สำคัญ",
    "ระบบต้องแยกข้อมูลบริการประชาชนออกจากข้อมูลสำนวนคดี และไม่เปิดเผยข้อมูลคดีหรือข้อมูลส่วนบุคคลแก่ อสม. เครือข่าย หรือบุคคลที่ไม่มีอำนาจหน้าที่",
]:
    add_number_item(doc, text, num_id=num_scope, under=True)

add_section_heading(doc, "6", "วิธีดำเนินการ")
num_steps = new_numbering_group(doc)
steps = [
    "แต่งตั้งคณะทำงานด้านบริหารโครงการ ด้านกฎหมาย ด้านเทคโนโลยีสารสนเทศ ด้านคุ้มครองข้อมูลส่วนบุคคล และตัวแทนผู้ใช้งาน พร้อมกำหนดผู้รับผิดชอบแต่ละงาน",
    "ประชุมคณะทำงานเพื่อวิเคราะห์กระบวนงาน กำหนดความต้องการ พจนานุกรมข้อมูล แบบฟอร์มทางกฎหมาย ระดับสิทธิผู้ใช้ และเกณฑ์รับมอบระบบ จำนวน 1 ครั้ง ผู้เข้าร่วม 20 คน",
    "จัดทำมาตรฐานวิธีปฏิบัติงาน (SOP) สำหรับการรวบรวมพยานหลักฐานดิจิทัล การคำนวณค่าแฮช การจัดทำ Chain of Custody การตรวจทานผลวิเคราะห์ และการส่งออกชุดหลักฐาน",
    "ประเมินความเสี่ยงด้านข้อมูลส่วนบุคคลและความมั่นคงปลอดภัย จัดทำรายการกิจกรรมประมวลผลข้อมูล กำหนดระยะเวลาเก็บรักษาและทำลายข้อมูล รวมถึงแผนรับมือเหตุละเมิดหรือระบบขัดข้อง",
    "ปรับปรุงระบบและทดสอบด้วยข้อมูลจำลองหรือข้อมูลที่ปกปิดตัวบุคคล ก่อนใช้ข้อมูลจริง พร้อมทดสอบการควบคุมสิทธิ Audit Log การสำรองและกู้คืนข้อมูล",
    "ทดสอบการยอมรับโดยผู้ใช้ (UAT) กับกรณีศึกษาหรือเรื่องร้องเรียนไม่น้อยกว่า 20 เรื่อง โดยคณะทำงานด้านกฎหมายตรวจสอบผลทุกกรณี และแก้ไขข้อบกพร่องระดับร้ายแรงก่อนเปิดใช้งาน",
    "จัดอบรมเชิงปฏิบัติการจำนวน 1 ครั้ง รวม 120 คน โดยแบ่งเนื้อหาเป็นหลักสูตรเจ้าหน้าที่ 90 คน และหลักสูตร อสม./เครือข่ายภาคประชาชน 30 คน",
    "เปิดใช้งานนำร่อง ติดตามสถิติการใช้งาน ระยะเวลาปฏิบัติงาน ความถูกต้องของผลวิเคราะห์ เหตุการณ์ด้านความปลอดภัย และข้อเสนอแนะของผู้ใช้ เป็นเวลาไม่น้อยกว่า 3 เดือน",
    "สรุปผลการดำเนินงานและเสนอผู้บริหารพิจารณาการขยายผลหรือปรับปรุงระบบในระยะต่อไป",
]
for text in steps:
    add_number_item(doc, text, num_id=num_steps, under=True)

add_section_heading(doc, "7", "สถานที่ดำเนินการ")
add_plain(doc, "ห้องประชุมสำนักงานสาธารณสุขจังหวัดศรีสะเกษ และระบบประชุมออนไลน์ของหน่วยงาน")
add_under(doc, "การทดสอบและใช้งานระบบให้ดำเนินการผ่านอุปกรณ์และเครือข่ายที่หน่วยงานอนุญาต โดยหลีกเลี่ยงการจัดเก็บพยานหลักฐานลงในอุปกรณ์ส่วนบุคคล")

add_section_heading(doc, "8", "ระยะเวลาดำเนินการ")
add_under(doc, "เดือนสิงหาคม พ.ศ. 2569 ถึงเดือนธันวาคม พ.ศ. 2569 โดยกิจกรรมที่มีการเบิกจ่ายจากงบประมาณประจำปี พ.ศ. 2569 ให้ดำเนินการภายในวันที่ 30 กันยายน พ.ศ. 2569 ส่วนการติดตามผลตั้งแต่เดือนตุลาคมถึงธันวาคม พ.ศ. 2569 ให้ใช้ทรัพยากรประจำของหน่วยงานและดำเนินการตามความเห็นชอบของฝ่ายการเงิน")
add_schedule_table(doc)

add_section_heading(doc, "9", "งบประมาณ")
add_plain(doc, "ค่าใช้จ่ายในการดำเนินงานเบิกจากงบประมาณโครงการพัฒนางานคุ้มครองผู้บริโภคด้านผลิตภัณฑ์สุขภาพ สำนักงานสาธารณสุขจังหวัดศรีสะเกษ รวมเป็นเงินทั้งสิ้น 22,600 บาท (สองหมื่นสองพันหกร้อยบาทถ้วน) โดยมีรายละเอียดดังนี้")
add_budget_table(doc)
add_under(doc, "การวิเคราะห์ ออกแบบ ปรับปรุง ทดสอบ และดูแลระบบดำเนินการโดยบุคลากรของหน่วยงานหรือทรัพยากรที่หน่วยงานมีอยู่เดิม โดยไม่เบิกค่าใช้จ่ายเพิ่มเติมจากโครงการนี้ หากมีค่าใช้บริการภายนอก ค่าโครงสร้างพื้นฐาน หรือค่าบำรุงรักษาเพิ่มเติม ให้เสนอขออนุมัติและดำเนินการตามระเบียบที่เกี่ยวข้องแยกต่างหาก")
add_plain(doc, "หมายเหตุ: ค่าใช้จ่ายทุกรายการสามารถถัวเฉลี่ยกันได้ตามที่จ่ายจริง โดยไม่เกินอัตราตามระเบียบของทางราชการ และอยู่ภายในวงเงินงบประมาณที่ได้รับอนุมัติ", first_line=0)

add_section_heading(doc, "10", "การประเมินผล")
add_subheading(doc, "10.1", "ตัวชี้วัดโครงการ")
num_kpi = new_numbering_group(doc)
for text, under in [
    ("มีระบบ LawiRisk-SSK Evidence Intelligence ที่สามารถจัดเก็บ สืบค้น วิเคราะห์ และแสดงความเชื่อมโยงของพยานหลักฐานดิจิทัลได้", False),
    ("กรณีทดสอบสำคัญผ่านเกณฑ์ร้อยละ 100 และไม่มีข้อบกพร่องหรือช่องโหว่ระดับร้ายแรงที่ยังไม่ได้แก้ไขก่อนเปิดใช้งาน", True),
    ("พยานหลักฐานในกรณีทดสอบมีแหล่งที่มา วันเวลา ผู้เก็บ ค่าแฮช และประวัติ Chain of Custody ครบถ้วนร้อยละ 100", True),
    ("ความแม่นยำและความครอบคลุมของการแสดงจุดเชื่อมโยงเป็นไปตามเกณฑ์ที่คณะทำงานกำหนดก่อนทดสอบ โดยเสนอเป้าหมายไม่น้อยกว่าร้อยละ 90 ในชุดข้อมูลที่ผู้เชี่ยวชาญรับรอง", True),
    ("ระยะเวลามัธยฐานในการสืบค้น วิเคราะห์ และจัดทำร่างเอกสารลดลงไม่น้อยกว่าร้อยละ 60 หรือเหลือไม่เกิน 2 ชั่วโมงต่อเรื่อง", True),
    ("ผู้เข้ารับการอบรมไม่น้อยกว่าร้อยละ 80 มีคะแนนหลังอบรมไม่น้อยกว่าร้อยละ 80", True),
    ("ผู้ได้รับสิทธิใช้งานอย่างน้อยร้อยละ 80 มีการเข้าใช้งานจริงภายใน 3 เดือนหลังเปิดใช้", True),
    ("ไม่มีการเปิดเผยข้อมูลคดีแก่ผู้ไม่มีอำนาจ และเหตุการณ์ด้านความปลอดภัยที่เกิดขึ้นได้รับการบันทึก ตรวจสอบ และดำเนินการตามแผนตอบสนองร้อยละ 100", True),
]:
    add_number_item(doc, text, num_id=num_kpi, under=under)

add_subheading(doc, "10.2", "เครื่องมือและวิธีการประเมินผล")
num_eval = new_numbering_group(doc)
for text, under in [
    ("แบบทดสอบก่อนและหลังการอบรมสำหรับผู้ใช้งานแต่ละบทบาท", True),
    ("แบบบันทึกเวลาและขั้นตอนการปฏิบัติงานก่อนและหลังใช้ระบบ", True),
    ("ชุดกรณีทดสอบที่คณะทำงานด้านกฎหมายรับรอง พร้อมแบบตรวจความถูกต้องและความครอบคลุมของจุดเชื่อมโยง", True),
    ("รายงาน UAT รายงานการทดสอบสิทธิผู้ใช้ รายงานช่องโหว่ และหลักฐานการแก้ไข", True),
    ("รายงาน Audit Log สถิติการเข้าใช้งาน รายงานเหตุผิดปกติ และผลการทดสอบการกู้คืนข้อมูล", True),
    ("แบบประเมินความพึงพอใจและข้อเสนอแนะของเจ้าหน้าที่ อสม. และเครือข่ายตามขอบเขตที่แต่ละกลุ่มได้รับอนุญาต", True),
]:
    add_number_item(doc, text, num_id=num_eval, under=under)

add_section_heading(doc, "11", "ผลที่คาดว่าจะได้รับ")
add_subheading(doc, "11.1", "เชิงผลผลิต (Output)")
num_output = new_numbering_group(doc)
for text, under in [
    ("ได้ระบบ LawiRisk-SSK Evidence Intelligence สำหรับสืบค้น วิเคราะห์ความเชื่อมโยง และบริหารจัดการพยานหลักฐานดิจิทัล", True),
    ("ได้มาตรฐานวิธีปฏิบัติงาน คู่มือผู้ใช้ คู่มือผู้ดูแลระบบ แบบฟอร์ม Chain of Custody และเกณฑ์ตรวจรับระบบ", True),
    ("มีเจ้าหน้าที่และเครือข่ายผ่านการอบรมตามบทบาทและขอบเขตสิทธิที่กำหนด", True),
]:
    add_number_item(doc, text, num_id=num_output, under=under)
add_subheading(doc, "11.2", "เชิงผลลัพธ์ (Outcome)")
num_outcome = new_numbering_group(doc)
for text, under in [
    ("กลุ่มงานคุ้มครองผู้บริโภคสามารถจัดการพยานหลักฐานของเรื่องร้องเรียนและเชื่อมโยงข้อมูลข้ามคดีได้รวดเร็วและตรวจสอบย้อนกลับได้", True),
    ("พนักงานเจ้าหน้าที่ลดเวลางานซ้ำซ้อนและมีข้อมูลประกอบการพิจารณาที่เป็นระบบ โดยยังคงเป็นผู้รับผิดชอบการวินิจฉัยข้อเท็จจริงและข้อกฎหมาย", True),
    ("ผู้บริโภคได้รับการคุ้มครองจากการเฝ้าระวังและบังคับใช้กฎหมายที่รวดเร็วขึ้น โดยข้อมูลส่วนบุคคลและข้อมูลคดีได้รับการป้องกันตามสิทธิและความจำเป็น", True),
]:
    add_number_item(doc, text, num_id=num_outcome, under=under)

add_section_heading(doc, "12", "การกำกับดูแลข้อมูล กฎหมาย และความมั่นคงปลอดภัย")
num_gov = new_numbering_group(doc)
for text in [
    "กำหนดให้สำนักงานสาธารณสุขจังหวัดศรีสะเกษเป็นเจ้าของระบบและผู้ควบคุมข้อมูลตามบทบาทหน้าที่ โดยมอบหมายผู้ดูแลระบบ ผู้รับผิดชอบข้อมูล ผู้ตรวจทานด้านกฎหมาย และผู้ประสานงานด้านคุ้มครองข้อมูลส่วนบุคคลเป็นลายลักษณ์อักษร",
    "จัดทำรายการกิจกรรมการประมวลผลข้อมูลส่วนบุคคล กำหนดวัตถุประสงค์และฐานการประมวลผล ประเภทข้อมูล ผู้รับข้อมูล ระยะเวลาเก็บรักษา มาตรการรักษาความมั่นคงปลอดภัย และช่องทางใช้สิทธิของเจ้าของข้อมูล",
    "ใช้หลักการให้สิทธิเท่าที่จำเป็นต่อหน้าที่ ทบทวนสิทธิเป็นระยะ และเพิกถอนสิทธิทันทีเมื่อพ้นหน้าที่หรือยุติการปฏิบัติงาน",
    "เก็บต้นฉบับพยานหลักฐานแยกจากข้อมูลที่ระบบสร้างขึ้น ห้ามเขียนทับต้นฉบับ และให้ตรวจสอบค่าแฮชก่อนส่งมอบหรือส่งออกหลักฐาน",
    "ผลวิเคราะห์ของระบบต้องแสดงแหล่งข้อมูลและข้อจำกัด ห้ามนำผลที่ระบบสร้างขึ้นไปใช้กล่าวหาหรือดำเนินคดีโดยไม่มีการตรวจสอบจากพนักงานเจ้าหน้าที่",
    "จัดทำแผนสำรองและกู้คืนข้อมูล แผนรับมือเหตุละเมิดข้อมูลหรือระบบถูกโจมตี และช่องทางรายงานเหตุแก่ผู้บริหาร ผู้รับผิดชอบความมั่นคงปลอดภัย และผู้เกี่ยวข้อง",
    "ทบทวนให้สอดคล้องกับกฎหมายว่าด้วยธุรกรรมทางอิเล็กทรอนิกส์ กฎหมายคุ้มครองข้อมูลส่วนบุคคล กฎหมายเฉพาะที่เกี่ยวข้อง และแนวทางธรรมาภิบาลปัญญาประดิษฐ์ที่หน่วยงานกำหนด",
]:
    add_number_item(doc, text, num_id=num_gov, under=True)

add_section_heading(doc, "13", "แผนบำรุงรักษาและความต่อเนื่อง")
num_maint = new_numbering_group(doc)
for text in [
    "มอบหมายผู้ดูแลระบบหลักและผู้ปฏิบัติหน้าที่แทน พร้อมกำหนดช่องทางแจ้งปัญหาและระยะเวลาตอบสนองตามระดับความรุนแรง",
    "สำรองข้อมูลตามรอบที่กำหนดและทดสอบการกู้คืนอย่างน้อยปีละ 2 ครั้ง หรือเมื่อมีการเปลี่ยนแปลงระบบที่สำคัญ",
    "ติดตามข้อผิดพลาด ช่องโหว่ และความถูกต้องของผลวิเคราะห์ พร้อมบันทึกเวอร์ชันและการเปลี่ยนแปลงของระบบ",
    "ทบทวนแบบฟอร์ม ฐานกฎหมาย แหล่งข้อมูล และสิทธิผู้ใช้อย่างน้อยปีละ 1 ครั้ง หรือเมื่อกฎหมายหรือกระบวนงานเปลี่ยนแปลง",
    "จัดทำแผนส่งออกข้อมูลและย้ายระบบเพื่อป้องกันการผูกติดกับผู้ให้บริการรายเดียว และให้หน่วยงานสามารถเข้าถึงข้อมูลและรหัสต้นฉบับที่จำเป็นต่อการดำเนินงานต่อเนื่อง",
]:
    add_number_item(doc, text, num_id=num_maint, under=True)

add_section_heading(doc, "14", "ผู้รับผิดชอบโครงการ")
add_plain(doc, "กลุ่มงานคุ้มครองผู้บริโภคและเภสัชสาธารณสุข สำนักงานสาธารณสุขจังหวัดศรีสะเกษ")
add_under(doc, "คณะทำงานกำกับดูแลระบบ LawiRisk-SSK Evidence Intelligence ซึ่งประกอบด้วยผู้แทนด้านกฎหมาย เทคโนโลยีสารสนเทศ คุ้มครองข้อมูลส่วนบุคคล และผู้ใช้งานระบบ")

add_section_heading(doc, "15", "เอกสารประกอบที่ต้องจัดทำก่อนเปิดใช้ข้อมูลจริง")
num_docs = new_numbering_group(doc)
for text in [
    "คำสั่งแต่งตั้งคณะทำงานและตารางกำหนดผู้รับผิดชอบ (RACI)",
    "ขอบเขตความต้องการและเกณฑ์ตรวจรับระบบ",
    "พจนานุกรมข้อมูลและตารางสิทธิผู้ใช้",
    "SOP การจัดการพยานหลักฐานดิจิทัลและแบบฟอร์ม Chain of Custody",
    "รายการกิจกรรมประมวลผลข้อมูลส่วนบุคคล ประกาศความเป็นส่วนตัว และตารางระยะเวลาเก็บรักษาข้อมูล",
    "รายงาน UAT รายงานทดสอบความมั่นคงปลอดภัย และหลักฐานการแก้ไขข้อบกพร่อง",
    "แผนสำรองและกู้คืนข้อมูล แผนรับมือเหตุผิดปกติ และแผนบำรุงรักษา",
    "คู่มือผู้ใช้ คู่มือผู้ดูแลระบบ และหลักฐานการฝึกอบรม",
]:
    add_number_item(doc, text, num_id=num_docs, under=True)

signature_heading = add_plain(doc, "การลงนาม", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=None, after=18, line=1.0)
signature_heading.paragraph_format.page_break_before = True
signatures = [
    ("ผู้เสนอโครงการ", "(นศ.ภ. จตุพร เรืองวิไลกฤตย์)", "นักศึกษาเภสัชศาสตร์ชั้นปีที่ 6"),
    ("ผู้เสนอโครงการ", "(นางสาวมัลลิกา สุพล)", "เภสัชกรชำนาญการ"),
    ("ผู้เห็นชอบโครงการ", "(นายไพฑูรย์ แก้วภมร)", "เภสัชกรเชี่ยวชาญ\n(รองนายแพทย์สาธารณสุขจังหวัดศรีสะเกษ)"),
    ("ผู้อนุมัติโครงการ", "(นายทนง วีระแสงพงษ์)", "นายแพทย์สาธารณสุขจังหวัดศรีสะเกษ\nปฏิบัติราชการแทนผู้ว่าราชการจังหวัดศรีสะเกษ"),
]
for role, name, title in signatures:
    add_plain(doc, f"(ลงชื่อ) ............................................................ {role}",
              align=WD_ALIGN_PARAGRAPH.CENTER, first_line=None, after=3, line=1.0)
    add_plain(doc, name, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=None, after=0, line=1.0)
    for line in title.split("\n"):
        add_plain(doc, line, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=None, after=0, line=1.0)
    add_plain(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=None, after=10, line=1.0)

# Document metadata and final paragraph controls.
doc.core_properties.title = "ร่างโครงการ LawiRisk-SSK Evidence Intelligence"
doc.core_properties.subject = "ฉบับขีดเส้นใต้ข้อความเสนอเพิ่มเติมและแก้ไข"
doc.core_properties.author = "สำนักงานสาธารณสุขจังหวัดศรีสะเกษ"
doc.core_properties.keywords = "LawiRisk-SSK, Evidence Intelligence, พยานหลักฐานดิจิทัล, Chain of Custody"

for p in doc.paragraphs:
    p.paragraph_format.widow_control = True

doc.save(OUT)
print(OUT)
